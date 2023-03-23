from docx import Document
from docx.shared import Pt

class CalculoRecibo():
    
    @classmethod
    def gerar(cls, valores_iniciais):
        valor = valores_iniciais['valor']
        valor_extenso = valores_iniciais['v_extenso']
        dt_inicio_contrato = valores_iniciais['data_inicio_contrato']
        dt_final_contrato = valores_iniciais['data_final_contrato']
        
        
        lista_recibos = []
        mes = 0
        for periodo in range(12):
            
            
            dados_recibo = {}
            dados_recibo['{VALOR}'] = valor
            dados_recibo['{VALOR_EXTENSO}'] = valor_extenso
            
            if periodo == 0:
                dados_recibo['mes'] = str(periodo + 1)
                dados_recibo['{VENCIMENTO}'] = cls.soma_mes(dt_inicio_contrato, 0)
                dados_recibo['{MES_REFER}'] = cls.mes_referencia(dados_recibo['{VENCIMENTO}'])
                dados_recibo['{PERIODO}'] = cls.periodo_ref(dados_recibo['{VENCIMENTO}'])
                
            else:
                mes = mes + 1
                dados_recibo['mes'] = str(mes + 1)
                
                dados_recibo['{VENCIMENTO}'] = cls.soma_mes(dt_inicio_contrato, mes)
                dados_recibo['{MES_REFER}'] = cls.mes_referencia(dados_recibo['{VENCIMENTO}'])
                dados_recibo['{PERIODO}'] = cls.periodo_ref(dados_recibo['{VENCIMENTO}'])
                
            dados_recibo['{CONTRATO_INICIO}'] = dt_inicio_contrato
            dados_recibo['{CONTRATO_FINAL}'] = dt_final_contrato
                
                
            
            lista_recibos.append(dados_recibo)
                
        return lista_recibos
                
    @classmethod
    def soma_mes(cls, data, numero_mes):
        
        split = data.split('/')
        mes = int(split[1])
        mes = mes + numero_mes
        
        if mes > 12:
            ano = int(split[2])
            ano = ano + 1
            mes = mes - 12
            nova_data = split[0] + '/' + str(mes) + '/' + str(ano)
        else:
            nova_data = split[0] + '/' + str(mes) + '/' + split[2]
            
        return nova_data
    
    @classmethod
    def mes_referencia(cls, data, novo_mes=None):
        
        split = data.split('/')
        
        mes = int(split[1])
        ano = int(split[2])
        
        if novo_mes != None:
            if mes > 12:
                mes = mes - 12
                ano = ano + 1
            m_referencia = str(mes) + '/' + str(ano)
        else:
            m_referencia = split[1] + '/' + split[2]
        
        return m_referencia
    
    @classmethod
    def periodo_ref(cls, data_vencimento):
        
        split = data_vencimento.split('/')
        dia = int(split[0]) - 1
        mes = int(split[1]) + 1
        ano = int(split[2])
        
        
        if mes > 12:
            mes = mes - 12
            ano = ano + 1
        
        data_ref_final = "{}/{}/{}".format(str(dia), str(mes), str(ano))
        
        return "{} à {}".format(data_vencimento, data_ref_final)
    
    @classmethod
    def mes_referencia(cls, data):
        
        split = data.split('/')
        
        m_referencia = split[1] + '/' + split[2]
        
        return m_referencia



class Recibos():
    
    def __init__(self, path_cliente, calculo_recibo):
        self.path_pasta_recibos = path_cliente + '/RECIBOS'
        self.path_cliente = path_cliente
        self.calculo_recibo = calculo_recibo
        self.arquivo_nome = self.nome_arquivo(calculo_recibo[0]['{CONTRATO_FINAL}'])
    
    def salvar(self):
    
        self.modelo_recibo.save(self.arquivo_nome)
        
        return self.arquivo_nome
    
    
    def _carregar_doc_modelo(self):
        path_recibo_modelo = self.path_cliente + '/RECIBOS/MODELO.docx'
        
        return Document(path_recibo_modelo)
        
    
    def criar_doc(self):
        self._criar_doc()
        
        return self.modelo_recibo
    
    def _gerar(self):
        self.modelo_recibo = self._carregar_doc_modelo()
        self.tabelas = self.modelo_recibo.tables
        
        
    def _criar_doc(self):
        self._gerar()
        
        contador_recibos = 0
        contador_titulos = 0
        
        for index, tabela in enumerate(self.tabelas):
            rows = tabela.rows
            
            for row in rows:
                for cell in row.cells:
                    
                    if cell.text == 'RECIBO DE ALUGUEL' and index > 0:
                        contador_titulos += 1
                    
                    if contador_titulos > 1:
                        contador_recibos += 1
                        contador_titulos = 0
                        
                    if contador_recibos < 12:
                        for paragrafo in cell.paragraphs:
                            novo_texto = self._substituir_valor(self.calculo_recibo[contador_recibos], paragrafo.text)
                            paragrafo.text = novo_texto
                            paragrafo.style.font.size = Pt(10.5)
                            
                            runs = paragrafo.runs
                            for run in runs:
                                run.bold = True
                            
                            
    
    def _substituir_valor(self, recibo, texto):
        for key in recibo.keys():
            texto = texto.replace(key, recibo[key])
        return texto
    
    def nome_arquivo(self, dt_final_contrato):
        split = dt_final_contrato.split('/')
        ano = split[-1]
        nome_arquivo = self.path_pasta_recibos+ "/{}.docx".format(ano)
        
        return nome_arquivo
